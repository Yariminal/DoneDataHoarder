"""
Document analyzer — extracts text from PDFs and Office files, then summarises with LLM.

Supported formats: PDF, DOCX, XLSX, TXT, CSV, and other text-based files.
Only sends a limited excerpt to the AI to avoid context-window issues.

For PDFs whose text extraction yields nothing (image-only / scanned docs /
design-heavy brochures like event menus), this module falls back to
rendering the first page as a JPEG and sending it to the *vision* model —
so we still get a meaningful name and tags instead of a blind filename guess.
"""
import io
from pathlib import Path
from typing import Optional

from datahoarder.analyzers.base import AnalysisResult, BaseAnalyzer, SYSTEM_PROMPT
from datahoarder.db.models import File

MAX_CHARS = 3000   # max text chars to send to AI
MAX_PAGES = 3      # max PDF pages to read

# Hard size cap on document *text* extraction (pdfplumber, openpyxl, etc.).
# Files larger than this skip text extraction entirely and the analyzer
# falls back to the rendered-page vision path (for PDFs) or to filename /
# folder context.
#
# Why: pdfplumber can take minutes / OOM on very large PDFs (the Solar
# Dekathlon test has a 452 MB PDF that hung the analyzer entirely),
# openpyxl / python-docx have similar pathological cases. We'd rather get
# a degraded analysis than no analysis at all.
MAX_DOC_SIZE_BYTES = 100 * 1024 * 1024  # 100 MB

# Size cap for *rendering* a PDF via pypdfium2. Rendering is per-page and
# lazy (Chromium PDFium streams pages rather than loading the whole file
# into memory), so it tolerates far larger files than text extraction:
# the 452 MB test PDF opens in 10 ms and renders page 0 in 120 ms with
# only +30 MB RSS. We keep a generous 2 GB ceiling as a safety rail for
# truly pathological files (backups accidentally named .pdf, etc.).
MAX_PDF_RENDER_BYTES = 2 * 1024 * 1024 * 1024  # 2 GB

# Number of pages to render + send to the vision model for rich PDFs.
# Matches MAX_PAGES on the text side for symmetry. Three pages is enough
# to cover cover + table-of-contents + first content page for most
# brochures / reports, at ~120 KB per JPEG → ~360 KB total upload.
MAX_VISION_PAGES = 3

DOC_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".odt",
    ".xlsx", ".xls", ".ods", ".csv",
    ".pptx", ".ppt", ".odp",
    ".txt", ".md", ".rtf",
    ".json", ".xml", ".yaml", ".yml",
    ".html", ".htm",
    ".ai",   # Adobe Illustrator (PDF-based — text extraction often works)
    ".mtl",  # Wavefront material library — plain text
}
DOC_MIMES = {
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.ms-excel",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-powerpoint",
    "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    "text/plain", "text/csv", "text/html",
    "application/json", "application/xml", "text/xml",
    "application/rtf", "text/rtf",
}


# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------

def _file_too_big(path: Path) -> bool:
    """True if *text* extracting this file is likely to OOM / hang."""
    try:
        return path.stat().st_size > MAX_DOC_SIZE_BYTES
    except OSError:
        return False


def _file_too_big_for_render(path: Path) -> bool:
    """True only for truly pathological PDFs (multi-GB). Rendering via
    pypdfium2 is lazy and page-scoped, so the 100 MB text-extraction
    gate is far too conservative for it."""
    try:
        return path.stat().st_size > MAX_PDF_RENDER_BYTES
    except OSError:
        return False


def _extract_pdf(path: Path) -> str:
    if _file_too_big(path):
        # Skip PDFs over 100 MB — pdfplumber's per-page parsing scales badly
        # on huge files (especially scanned docs and embedded-image-heavy PDFs).
        # The analyzer will fall back to filename/folder inference.
        return ""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages[:MAX_PAGES]:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n\n".join(text_parts)
    except ImportError:
        # Fallback: try PyPDF2 or just give up
        return ""
    except Exception:
        return ""


# Vision-fallback render settings. Scale 2.0 ~= 144 DPI for typical PDFs —
# enough resolution for the vision model to read headings and small text
# while keeping the image small. We then downsize to PDF_VISION_MAX_SIDE
# before JPEG-encoding so the upload to the model stays cheap.
PDF_RENDER_SCALE = 2.0
PDF_VISION_MAX_SIDE = 1024
PDF_VISION_JPEG_QUALITY = 85


def _render_pdf_pages_as_jpegs(
    path: Path,
    max_pages: int = MAX_VISION_PAGES,
) -> list[bytes]:
    """
    Render up to `max_pages` pages of a PDF as JPEG byte blobs suitable
    for a vision model. Returns an empty list on any failure (missing
    library, encrypted / corrupt PDF, zero pages, over-size file) so the
    caller can gracefully degrade. Never raises.

    Size gate is MAX_PDF_RENDER_BYTES (2 GB), much higher than the
    text-extraction cap — pypdfium2 handles a 452 MB PDF in 120 ms / +30
    MB RSS because it streams one page at a time. The text-extraction
    cap is specifically about pdfplumber's appetite for the whole file.

    Returning a list (rather than the old single-page helper) lets the
    caller send the cover + first 2 content pages together via the
    `images_list` kwarg on the vision client — much richer context for
    content-dense PDFs like multi-page brochures, reports, and project
    documents (see the Solar Dekathlon 15.12.pdf case: 452 MB, 33 pages
    of text and 3D renderings that need more than a cover to name well).
    """
    if _file_too_big_for_render(path):
        return []
    try:
        import pypdfium2 as pdfium
        from PIL import Image as PilImage
    except ImportError:
        return []

    try:
        pdf = pdfium.PdfDocument(str(path))
    except Exception:
        return []

    jpegs: list[bytes] = []
    try:
        total = len(pdf)
        if total == 0:
            return []
        # Render the first N pages — simpler and more predictable than
        # sampling e.g. first/middle/last. Can revisit if vision quality
        # plateaus. Page indices are 0-based.
        for idx in range(min(max_pages, total)):
            try:
                page = pdf[idx]
                pil = page.render(scale=PDF_RENDER_SCALE).to_pil()
            except Exception:
                # One bad page shouldn't nuke the whole render — skip it
                # and try the next. If every page fails we return [].
                continue
            try:
                pil = pil.convert("RGB")
                w, h = pil.size
                if max(w, h) > PDF_VISION_MAX_SIDE:
                    ratio = PDF_VISION_MAX_SIDE / max(w, h)
                    pil = pil.resize(
                        (int(w * ratio), int(h * ratio)),
                        PilImage.LANCZOS,
                    )
                buf = io.BytesIO()
                pil.save(buf, format="JPEG", quality=PDF_VISION_JPEG_QUALITY)
                jpegs.append(buf.getvalue())
            except Exception:
                continue
    finally:
        # PdfDocument holds a native PDFium handle; release it promptly.
        try:
            pdf.close()
        except Exception:
            pass

    return jpegs


def _render_pdf_first_page_as_jpeg(path: Path) -> bytes | None:
    """Thin compatibility wrapper — returns just the first rendered
    page's JPEG bytes, or None. Retained so existing callers / tests
    that expect the single-page helper keep working."""
    pages = _render_pdf_pages_as_jpegs(path, max_pages=1)
    return pages[0] if pages else None


def _extract_docx(path: Path) -> str:
    if _file_too_big(path):
        return ""
    try:
        from docx import Document
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except ImportError:
        return ""
    except Exception:
        return ""


def _extract_xlsx(path: Path) -> str:
    if _file_too_big(path):
        return ""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        rows = []
        for sheet in wb.worksheets[:2]:           # first 2 sheets
            for row in sheet.iter_rows(max_row=20, values_only=True):
                cell_vals = [str(c) for c in row if c is not None]
                if cell_vals:
                    rows.append(", ".join(cell_vals))
            if len(rows) > 30:
                break
        wb.close()
        return "\n".join(rows)
    except ImportError:
        return ""
    except Exception:
        return ""


def _extract_text(path: Path) -> str:
    """Read plain-text files with encoding fallbacks."""
    if _file_too_big(path):
        # 1 GB log files / massive CSV exports — skip rather than load into RAM.
        # We only need MAX_CHARS worth anyway; future improvement would be a
        # streaming read of the first MAX_CHARS bytes, but for now skip.
        return ""
    for enc in ("utf-8", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=enc)
        except (UnicodeDecodeError, LookupError):
            continue
        except OSError:
            return ""
    return ""


def extract_text(path: Path, mime_type: Optional[str] = None) -> str:
    """Extract readable text from a document, up to MAX_CHARS."""
    ext = path.suffix.lower()
    mime = mime_type or ""

    if ext == ".pdf" or "pdf" in mime:
        text = _extract_pdf(path)
    elif ext in (".docx", ".doc") or "word" in mime:
        text = _extract_docx(path)
    elif ext in (".xlsx", ".xls") or "spreadsheet" in mime or "excel" in mime:
        text = _extract_xlsx(path)
    elif ext in (".txt", ".md", ".csv", ".json", ".xml", ".yaml", ".yml",
                  ".html", ".htm", ".rtf"):
        text = _extract_text(path)
    else:
        text = _extract_text(path)  # try anyway

    return text[:MAX_CHARS].strip()


# ---------------------------------------------------------------------------
# Analyzer class
# ---------------------------------------------------------------------------

PDF_VISION_PROMPT = """\
You are analyzing a PDF. Its first {page_count} page(s) have been
rendered and attached as images. The PDF's extractable text was empty
or too short to be useful — it's most likely a scanned document, a
design-heavy brochure, a flyer, a menu, a poster, a certificate, a
report with embedded 3D renderings, or a graphic cover. Describe what
you see (layout, imagery, visible headings, logos, colours, subject
matter) and propose a meaningful filename based on that.

If multiple pages are attached, treat them as an ordered excerpt — the
first is usually a cover / title, later pages show actual content.
Synthesise across all pages to produce a single description and name
rather than summarising each page separately.

Context about the file:
{context}

Return a JSON object with these fields:
{{
  "description": "1-2 sentences describing what the document is about, synthesised across the attached page(s)",
  "suggested_name": "meaningful filename stem. Rules: (1) Describe the document's actual content, do NOT repeat the containing folder name. (2) Preserve specific proper nouns (organisation, client, event name, project name) only if visible and uniquely identifying. (3) Translate to English if not already. (4) No extension, no date prefix, use_underscores, max 60 chars",
  "tags": ["tag1", "tag2", ...]  // 4-8 specific, lowercase, underscore_separated tags describing concrete visible elements (subject, document subtype, visible colour scheme, visible motifs/patterns, disciplines like "architecture" or "structural_engineering" if evident). Skip generic words like "document", "pdf", "page".,
  "document_type": "one of: invoice, receipt, contract, report, letter, cv_resume, photo, presentation, spreadsheet, notes, form, certificate, manual, menu, flyer, brochure, poster, cover, other",
  "detected_date": "YYYY-MM-DD only if a specific date is clearly visible on any rendered page — NOT inferred from the filename or folder. Return null if no explicit date is visible.",
  "language": "ISO 639-1 language code of any visible text (e.g. en, he, fr). If the page is purely graphical with no text, return null.",
  "confidence": 0.0-1.0
}}

For suggested_name: reflect what the document ACTUALLY shows. Examples:
- "event_menu_cactus_pattern" for a menu card with cactus illustrations
- "architecture_project_report_renderings" for a multi-page project report with 3D renderings
- "award_certificate_first_place" for a visible certificate layout
"""


DOC_PROMPT = """\
You are analyzing a document to help rename and categorize it in a personal file archive.

Context about the file:
{context}

Extracted text (first {max_chars} characters):
---
{text}
---

Based on the filename, folder context, and document content, return a JSON object:
{{
  "description": "1-2 sentences describing what this document is about",
  "suggested_name": "meaningful filename stem. Rules: (1) Describe the document's actual content/purpose — do NOT repeat the containing folder name. (2) Preserve specific proper nouns (client names, organizations) only if they uniquely identify this document. (3) Translate to English if not already. (4) No extension, no date prefix, use_underscores, max 60 chars",
  "tags": ["tag1", "tag2", ...]  // 4-8 specific, lowercase tags. RULES: (a) each tag must add information NOT already implied by the filename or folder name; (b) NO generic words like "document", "file", "text", "content"; (c) prefer concrete entities (organisation, client, project, topic, document subtype) over abstract categories; (d) no duplicates or near-duplicates; (e) use_underscores; (f) skip the tag rather than guessing if unsure,
  "document_type": "one of: invoice, receipt, contract, report, letter, cv_resume, photo, presentation, spreadsheet, notes, form, certificate, manual, other",
  "detected_date": "YYYY-MM-DD only if a specific date is explicitly written in the document text (e.g. '14 March 2021', 'Date: 2021-03-14') — NOT inferred from the folder name or filename. Return null if no explicit date is found.",
  "language": "ISO 639-1 language code (e.g. en, he, fr)",
  "confidence": 0.0-1.0
}}

For suggested_name: reflect the actual content.
Examples:
- "invoice_amazon_order_123" not "invoice"
- "employment_contract_2021" not "contract"
- "project_proposal_client_name" not "proposal"

If the text is empty or unreadable, use the filename and folder context to make your best guess.
"""


class DocumentAnalyzer(BaseAnalyzer):
    def __init__(self, ai_client):
        self._client = ai_client

    def can_handle(self, mime_type: str, extension: str) -> bool:
        if mime_type and mime_type in DOC_MIMES:
            return True
        return extension.lower() in DOC_EXTENSIONS

    def analyze(self, file_rec: File, context: str) -> AnalysisResult:
        path = Path(file_rec.path)
        text = extract_text(path, file_rec.mime_type)
        ext = path.suffix.lower()

        # Vision fallback: PDFs that yield near-zero extractable text are
        # almost always image-based (scans, menus, flyers, design covers)
        # OR over-large files where pdfplumber refused to parse (like the
        # 452 MB Solar Dekathlon 15.12.pdf report with 33 pages of text
        # and 3D renderings). Instead of letting the text-only LLM guess
        # from the filename, we rasterise up to MAX_VISION_PAGES pages and
        # ask the vision model what it sees — a dramatically better
        # signal for naming. _render returns [] on any failure (missing
        # pypdfium2, encrypted / corrupt PDF, render OOM, etc.).
        is_text_empty = not text or len(text.strip()) < 20
        if is_text_empty and ext == ".pdf":
            pdf_pages = _render_pdf_pages_as_jpegs(path)
            if pdf_pages:
                vision_result = self._analyze_rendered_pdf(
                    file_rec, context, pdf_pages
                )
                if vision_result is not None:
                    return vision_result
            # Fall through to the text-only path if rendering or vision
            # inference failed — the filename-only guess is still better
            # than no result at all.

        prompt = DOC_PROMPT.format(
            context=context,
            text=text or "(no readable text extracted)",
            max_chars=MAX_CHARS,
        )

        try:
            data = self._client.generate_json(prompt, system=SYSTEM_PROMPT)
        except Exception as exc:
            return AnalysisResult(
                description=f"AI inference failed: {exc}",
                confidence=0.0,
            )

        result = AnalysisResult.from_ai_response(data)
        doc_type = data.get("document_type", "")
        if doc_type and doc_type not in result.tags:
            result.tags.insert(0, doc_type)
        lang = data.get("language", "")
        if lang and lang != "en":
            result.tags.append(f"lang:{lang}")

        # If text extraction yielded nothing meaningful, the LLM was guessing
        # from filename + folder context only. Mark accordingly so the
        # description is prefixed [UNVERIFIED ...] and confidence is capped.
        if is_text_empty:
            result.content_available = False
            result.confidence = min(result.confidence, 0.4)

        return result

    def _analyze_rendered_pdf(
        self,
        file_rec: File,
        context: str,
        pages: list[bytes],
    ) -> AnalysisResult | None:
        """
        Vision path for text-empty PDFs. Sends up to MAX_VISION_PAGES
        rendered pages to the vision model and returns a full
        AnalysisResult. Returns None only if the vision call itself
        raised — caller falls back to the text-only (filename-guess)
        path in that case.

        For single-page inputs we use `image_bytes=` (cheaper / simpler
        path in both the Ollama and Gemini clients). For multi-page we
        use `images_list=` and the prompt tells the model to synthesise
        across pages rather than describe each.

        NOTE: `content_available=True` because the model really did see
        the pages — unlike the filename-only guess, this isn't
        [UNVERIFIED].
        """
        if not pages:
            return None
        prompt = PDF_VISION_PROMPT.format(
            context=context, page_count=len(pages)
        )
        kwargs: dict = {"system": SYSTEM_PROMPT}
        if len(pages) == 1:
            kwargs["image_bytes"] = pages[0]
        else:
            kwargs["images_list"] = pages

        try:
            data = self._client.generate_json(prompt, **kwargs)
        except Exception:
            # If the model doesn't support vision (or fails), let the
            # caller try the text-only path rather than returning an error.
            return None

        result = AnalysisResult.from_ai_response(data)
        doc_type = data.get("document_type", "")
        if doc_type and doc_type not in result.tags:
            result.tags.insert(0, doc_type)
        lang = data.get("language")
        if lang and lang != "en":
            result.tags.append(f"lang:{lang}")
        # Add a tag so the user / organizer can tell this PDF was
        # analyzed via rendered-page vision — useful for "why is this
        # tagged?" debugging and for deliberately routing similar files
        # later.
        if "image_only_pdf" not in result.tags:
            result.tags.append("image_only_pdf")
        return result
